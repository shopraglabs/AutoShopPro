#!/usr/bin/env python3
"""
update_build_log.py — Append a new session entry to AutoShopPro_Build_Log.docx
Usage: python3 docs/update_build_log.py

Reads session data from the SESSION dict below, then inserts the entry
directly after the "Build Sessions" heading in the document.

Requires: pip3 install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DOCX_PATH = os.path.join(os.path.dirname(__file__), 'AutoShopPro_Build_Log.docx')

# ── Fill this in at the end of each session ───────────────────────────────────
SESSION = {
    'heading':  'Session 22 — v0.13.0 Sharp Eye',
    'date':     'April 4, 2026',
    'version':  'v0.13.0 Sharp Eye',
    'number':   '22',
    'built': [
        'Search on estimates list — CupertinoSearchTextField added below nav bar; filters by estimate number (EST-XXXX), customer name, and vehicle (year/make/model); EstimateListScreen converted to ConsumerStatefulWidget',
        'Search on repair orders list — CupertinoSearchTextField added above filter pills; status filter and search apply simultaneously; empty states updated to reflect search vs filter context',
        'Default markup button on Add Part form — "Apply shop default markup" row (counterclockwise arrow icon) added between Markup fields and Unit List; reads defaultPartsMarkup from ShopSettings and recalculates sell price via existing _recalcFromCostAndPercent()',
    ],
    'bugs': [
        'Parts sub-header showed labor description instead of labor name — _LineItemSection used labor.description for group label; changed to labor.laborName ?? labor.description so "Brake job - Rear" shows instead of the full description sentence',
        'Parts sub-header overflowed (overflowed by ~109 pixels) — Text widget in Row had no width constraint; wrapped in Expanded with overflow: TextOverflow.ellipsis',
        'Service template save got stuck — _save() had no try/catch; any DB error left _saving=true with spinner showing forever; wrapped DB writes in try/catch and reset _saving=false on error; switched Navigator.pop to context.pop() for go_router correctness',
    ],
    'files_added': [],
    'files_modified': [
        'lib/features/estimates/estimate_list_screen.dart — converted ConsumerWidget to ConsumerStatefulWidget; added _search state; added CupertinoSearchTextField; filter logic by number/customer/vehicle; empty state handles both no-data and no-results cases',
        'lib/features/repair_orders/ro_list_screen.dart — added _search state; CupertinoSearchTextField above filter pills; combined status+search filtering; empty state text accounts for active search',
        'lib/features/estimates/line_item_form_screen.dart — added _applyDefaultMarkup() async method; added "Apply shop default markup" GestureDetector row between _MarkupRow and Unit List field',
        'lib/features/estimates/estimate_detail_screen.dart — _LineItemSection group label changed from labor.description to labor.laborName ?? labor.description; sub-header Text wrapped in Expanded with ellipsis overflow',
        'lib/features/service_templates/service_template_form_screen.dart — _save() wrapped in try/catch; _saving reset to false on error; error dialog shown; Navigator.pop replaced with context.pop()',
        'lib/features/invoices/invoice_detail_screen.dart — tax rate prefers taxRateBps; showInvoiceActions passes shop contact',
        'lib/features/invoices/invoice_service.dart — shopAddress/Phone/Email params throughout; PDF headers show contact; _deleteTempAfterDelay() helper cleans up temp files',
        'lib/features/estimates/estimate_detail_screen.dart — _convert() snapshots taxRateBps at RO creation',
        'lib/features/estimates/line_item_form_screen.dart — _pickFromCatalog uses _applyMarkupTier+_recalcFromCostAndPercent; _computeSellPriceCents() helper; template parts use fresh sell price; negative price allowed; $999,999.99 cap',
        'lib/features/settings/settings_screen.dart — shopAddress/Phone/Email controllers; _showAddRuleDialog auto-fills From cost as read-only; markup display uses fromCents()',
        'lib/features/service_templates/service_template_form_screen.dart — qtyController added to _LinkedPart typedef; qty field shown on each linked part row; qty saved/loaded from DB',
        'lib/features/dashboard/dashboard_screen.dart — "Revenue" → "Revenue (excl. tax)" labels',
        'lib/features/dashboard/dashboard_provider.dart — ARO uses countedROsForARO denominator',
        'docs/module_plan.md — v0.13.0 Sharp Eye section added',
        'README.md — v0.13.0 entry added at top of Project Status',
        'CLAUDE.md — version bumped to v0.13.0 Sharp Eye; estimates and ROs listed as searchable; parts display updated',
    ],
    'next': 'Continue polish: payment flows, more dashboard KPIs, or Phase 3 remaining items (#23 dbProvider to core, #24 taxAmountCents on invoice, #25 GDPR export).',
}
# ──────────────────────────────────────────────────────────────────────────────


def add_bold_label(para, label):
    run = para.add_run(label)
    run.bold = True


def add_kv_row(doc, label, value):
    """Add a label: value line (bold label, normal value)."""
    para = doc.add_paragraph(style='Normal')
    add_bold_label(para, f'{label}: ')
    para.add_run(value)


def add_section(doc, label, items):
    """Add a bold label followed by bulleted items."""
    if not items:
        return
    para = doc.add_paragraph(style='Normal')
    add_bold_label(para, label)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def insert_session(session):
    doc = Document(DOCX_PATH)

    # Find the paragraph index of "Build Sessions" heading
    insert_idx = None
    for i, para in enumerate(doc.paragraphs):
        if 'Build Sessions' in para.text:
            insert_idx = i
            break

    if insert_idx is None:
        print("ERROR: Could not find 'Build Sessions' heading in document.")
        return

    # python-docx can't insert paragraphs at an arbitrary position directly,
    # so we use the underlying XML to splice in after the target paragraph.
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    body = doc.element.body
    all_paras = body.findall(f'.//{qn("w:p")}')
    target_para_el = doc.paragraphs[insert_idx]._element

    def make_para(text, bold=False, style_id=None):
        p = OxmlElement('w:p')
        if style_id:
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), style_id)
            pPr.append(pStyle)
            p.append(pPr)
        r = OxmlElement('w:r')
        if bold:
            rPr = OxmlElement('w:rPr')
            b = OxmlElement('w:b')
            rPr.append(b)
            r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
        return p

    def make_bullet(text):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'ListBullet')
        pPr.append(pStyle)
        p.append(pPr)
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
        return p

    # Build the block of new XML elements (in reverse order, each inserted after target)
    blocks = []

    # Next session
    blocks.append(make_para('Next session'))
    blocks.append(make_para(session['next']))

    # Files modified
    if session['files_modified']:
        blocks.append(make_para('Files modified', bold=True))
        for f in session['files_modified']:
            blocks.append(make_bullet(f))

    # Files added
    if session['files_added']:
        blocks.append(make_para('Files added', bold=True))
        for f in session['files_added']:
            blocks.append(make_bullet(f))

    # Bugs
    if session['bugs']:
        blocks.append(make_para('Bugs fixed', bold=True))
        for b in session['bugs']:
            blocks.append(make_bullet(b))

    # What was built
    if session['built']:
        blocks.append(make_para('What was built', bold=True))
        for item in session['built']:
            blocks.append(make_bullet(item))

    # Metadata rows
    blocks.append(make_para(f"Session: {session['number']}"))
    blocks.append(make_para(f"Version: {session['version']}"))
    blocks.append(make_para(f"Date: {session['date']}"))

    # Section heading
    blocks.append(make_para(session['heading'], style_id='Heading2'))

    # Blank line before heading
    blocks.append(make_para(''))

    # Insert all blocks after the "Build Sessions" heading (reversed so order is correct)
    ref = target_para_el
    for block in blocks:
        ref.addnext(block)

    doc.save(DOCX_PATH)
    print(f"✅ Session entry '{session['heading']}' added to {DOCX_PATH}")


def read_build_log():
    """Print the build log contents to the terminal."""
    doc = Document(DOCX_PATH)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ''
        if 'Heading 1' in style:
            print(f'\n=== {text} ===')
        elif 'Heading 2' in style:
            print(f'\n--- {text} ---')
        elif 'List' in style:
            print(f'  • {text}')
        else:
            print(text)


if __name__ == '__main__':
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == 'read':
        read_build_log()
    else:
        insert_session(SESSION)
